"""
Certify Intel v7.0 - Knowledge Base RAG Pipeline
=================================================

Document ingestion, chunking, embedding, and retrieval-augmented generation.

Features:
- Multi-format document processing (PDF, DOCX, TXT, MD)
- Semantic chunking with overlap
- Batch embedding with OpenAI text-embedding-3-small
- Vector search via pgvector
- Citation tracking for all retrieved content
- Deduplication via content hashing

Usage:
    from knowledge_base import KnowledgeBase

    kb = KnowledgeBase()

    # Ingest a document
    result = await kb.ingest_document(
        file_path="/path/to/document.pdf",
        metadata={"competitor": "Epic Systems"}
    )

    # Search for relevant content
    results = await kb.search(
        query="What is Epic's pricing strategy?",
        limit=5
    )
"""

import os
import logging
import hashlib
from typing import Dict, List, Any, Optional
from dataclasses import dataclass, field
from datetime import datetime
from uuid import uuid4
import asyncio

logger = logging.getLogger(__name__)

USE_LOCAL_EMBEDDINGS = os.getenv("USE_LOCAL_EMBEDDINGS", "false").lower() == "true"


@dataclass
class DocumentChunk:
    """A chunk of a document."""
    chunk_id: str
    document_id: str
    chunk_index: int
    content: str
    token_count: int
    metadata: Dict[str, Any] = field(default_factory=dict)


@dataclass
class SearchResult:
    """A search result with similarity score."""
    chunk_id: str
    document_id: str
    content: str
    metadata: Dict[str, Any]
    similarity: float


@dataclass
class Document:
    """A document in the knowledge base."""
    document_id: str
    filename: str
    file_type: str
    content_hash: str
    chunk_count: int
    uploaded_by: Optional[str]
    metadata: Dict[str, Any]
    created_at: datetime


class KnowledgeBase:
    """
    Knowledge Base with RAG capabilities.

    Handles document ingestion, embedding, and retrieval.
    """

    # Supported file types
    SUPPORTED_TYPES = {
        ".pdf": "pdf",
        ".docx": "docx",
        ".doc": "docx",
        ".txt": "text",
        ".md": "markdown",
        ".html": "html"
    }

    def __init__(
        self,
        vector_store=None,
        embedding_model: str = "text-embedding-3-small",
        chunk_size: int = 500,
        chunk_overlap: int = 50
    ):
        self.vector_store = vector_store
        self.embedding_model = embedding_model
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self._openai_client = None

    async def _get_openai_client(self):
        """Lazy-load OpenAI client."""
        if self._openai_client is None:
            from openai import AsyncOpenAI
            self._openai_client = AsyncOpenAI()
        return self._openai_client

    async def ingest_document(
        self,
        file_path: str,
        uploaded_by: Optional[str] = None,
        metadata: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """
        Ingest a document into the knowledge base.

        Args:
            file_path: Path to the document
            uploaded_by: User ID who uploaded
            metadata: Additional metadata

        Returns:
            Dictionary with ingestion result
        """
        metadata = metadata or {}
        start_time = datetime.utcnow()

        try:
            # Step 1: Validate file
            if not os.path.exists(file_path):
                return {"status": "error", "error": f"File not found: {file_path}"}

            file_ext = os.path.splitext(file_path)[1].lower()
            if file_ext not in self.SUPPORTED_TYPES:
                return {"status": "error", "error": f"Unsupported file type: {file_ext}"}

            file_type = self.SUPPORTED_TYPES[file_ext]

            # Step 2: Extract text content
            content = await self._extract_content(file_path, file_type)

            if not content or len(content.strip()) < 10:
                return {"status": "error", "error": "No content extracted from document"}

            # Step 3: Check for duplicates
            content_hash = hashlib.sha256(content.encode()).hexdigest()

            existing = await self._check_duplicate(content_hash)
            if existing:
                return {
                    "status": "duplicate",
                    "existing_doc_id": existing["document_id"],
                    "message": "Document already exists in knowledge base"
                }

            # Step 4: Chunk the content
            chunks = self._chunk_content(content, file_path)

            if not chunks:
                return {"status": "error", "error": "Failed to chunk document"}

            # Step 5: Create document record
            doc_id = str(uuid4())
            document = Document(
                document_id=doc_id,
                filename=os.path.basename(file_path),
                file_type=file_type,
                content_hash=content_hash,
                chunk_count=len(chunks),
                uploaded_by=uploaded_by,
                metadata=metadata,
                created_at=datetime.utcnow()
            )

            # Store document record
            await self._store_document_record(document)

            # Step 6: Embed and store chunks in vector store
            embeddings = await self._batch_embed([chunk.content for chunk in chunks])

            if self.vector_store:
                import numpy as np
                from vector_store import DocumentChunk as VectorDocumentChunk

                # First insert the document record (returns existing ID if duplicate)
                actual_doc_id = await self.vector_store.insert_document(
                    document_id=doc_id,
                    filename=document.filename,
                    file_type=document.file_type,
                    content_hash=content_hash,
                    uploaded_by=uploaded_by or "system",
                    file_size_bytes=os.path.getsize(file_path),
                    metadata=metadata
                )

                # Use the actual document ID (might be existing if duplicate)
                doc_id = actual_doc_id

                # Then insert chunks with embeddings
                vector_chunks = [
                    VectorDocumentChunk(
                        chunk_index=chunk.chunk_index,
                        content=chunk.content,
                        embedding=np.array(emb),
                        token_count=chunk.token_count,
                        metadata={**chunk.metadata, **metadata}
                    )
                    for chunk, emb in zip(chunks, embeddings)
                ]

                await self.vector_store.batch_insert_chunks(
                    document_id=doc_id,
                    chunks=vector_chunks
                )

            processing_time = (datetime.utcnow() - start_time).total_seconds() * 1000

            return {
                "status": "success",
                "document_id": doc_id,
                "filename": document.filename,
                "chunks_created": len(chunks),
                "file_size_kb": os.path.getsize(file_path) / 1024,
                "processing_time_ms": processing_time
            }

        except Exception as e:
            logger.error(f"Document ingestion failed: {e}", exc_info=True)
            return {"status": "error", "error": str(e)}

    async def _extract_content(self, file_path: str, file_type: str) -> str:
        """Extract text content from a document."""
        try:
            if file_type == "pdf":
                return await self._extract_pdf(file_path)
            elif file_type == "docx":
                return await self._extract_docx(file_path)
            elif file_type in ["text", "markdown"]:
                with open(file_path, "r", encoding="utf-8") as f:
                    return f.read()
            elif file_type == "html":
                return await self._extract_html(file_path)
            else:
                return ""
        except Exception as e:
            logger.error(f"Content extraction failed: {e}")
            return ""

    async def _extract_pdf(self, file_path: str) -> str:
        """Extract text from PDF."""
        try:
            # Try unstructured first
            from unstructured.partition.pdf import partition_pdf

            elements = partition_pdf(file_path)
            return "\n\n".join(str(el) for el in elements)

        except ImportError:
            # Fallback to PyPDF2
            try:
                import PyPDF2

                with open(file_path, "rb") as f:
                    reader = PyPDF2.PdfReader(f)
                    text = []
                    for page in reader.pages:
                        text.append(page.extract_text() or "")
                    return "\n\n".join(text)

            except ImportError:
                logger.warning("No PDF extraction library available")
                return ""

    async def _extract_docx(self, file_path: str) -> str:
        """Extract text from DOCX."""
        try:
            from unstructured.partition.docx import partition_docx

            elements = partition_docx(file_path)
            return "\n\n".join(str(el) for el in elements)

        except ImportError:
            try:
                import docx

                doc = docx.Document(file_path)
                return "\n\n".join(p.text for p in doc.paragraphs)

            except ImportError:
                logger.warning("No DOCX extraction library available")
                return ""

    async def _extract_html(self, file_path: str) -> str:
        """Extract text from HTML."""
        try:
            from bs4 import BeautifulSoup

            with open(file_path, "r", encoding="utf-8") as f:
                soup = BeautifulSoup(f.read(), "html.parser")
                # Remove script and style elements
                for script in soup(["script", "style"]):
                    script.decompose()
                return soup.get_text(separator="\n")

        except ImportError:
            logger.warning("BeautifulSoup not available")
            return ""

    def _count_tokens(self, text: str) -> int:
        """Count tokens using tiktoken (accurate) or fallback to estimation."""
        try:
            import tiktoken
            enc = tiktoken.encoding_for_model("text-embedding-3-small")
            return len(enc.encode(text))
        except ImportError:
            # Fallback: ~4 characters per token
            return len(text) // 4

    def _split_into_sections(self, content: str) -> List[dict]:
        """
        Split content into semantic sections based on headers and structure.

        Returns list of dicts: [{"header": "...", "content": "...", "level": int}]
        """
        import re

        # Clean up PDF extraction artifacts (extra spaces)
        content = re.sub(r'  +', ' ', content)

        sections = []

        # Patterns for section headers (markdown and plain text)
        # More restrictive to avoid false positives
        header_patterns = [
            (r'^#{1,6}\s+(.+)$', 'markdown'),                    # # Header
            (r'^([A-Z][A-Z\s]{10,60})$', 'allcaps'),             # ALL CAPS HEADER (min 10 chars)
            (r'^(\d+\.\s+[A-Z][A-Za-z\s]{10,60})$', 'numbered'),  # 1. Header
        ]

        # Split on double newlines first (paragraph boundaries)
        paragraphs = re.split(r'\n\s*\n', content)

        current_section = {"header": "Document", "content": [], "level": 0}

        for para in paragraphs:
            para = para.strip()
            if not para:
                continue

            # Check if paragraph starts with a header
            first_line = para.split('\n')[0].strip()
            is_header = False

            for pattern, header_type in header_patterns:
                match = re.match(pattern, first_line)
                if match:
                    # Save previous section if it has substantial content
                    if current_section["content"]:
                        content_text = '\n\n'.join(current_section["content"]).strip()
                        if len(content_text) > 50:  # Only save if > 50 chars
                            sections.append({
                                "header": current_section["header"],
                                "content": content_text,
                                "level": current_section["level"]
                            })

                    # Start new section
                    header_text = match.group(1).strip()
                    rest_of_para = '\n'.join(para.split('\n')[1:]).strip()
                    level = 1 if header_type == 'markdown' else 2
                    current_section = {"header": header_text, "content": [], "level": level}

                    if rest_of_para:
                        current_section["content"].append(rest_of_para)

                    is_header = True
                    break

            if not is_header:
                current_section["content"].append(para)

        # Don't forget the last section
        if current_section["content"]:
            content_text = '\n\n'.join(current_section["content"]).strip()
            if len(content_text) > 50:
                sections.append({
                    "header": current_section["header"],
                    "content": content_text,
                    "level": current_section["level"]
                })

        # If no sections found or all content in one place, return as single section
        if not sections:
            sections = [{"header": "Content", "content": content, "level": 0}]

        return sections

    def _chunk_content(
        self,
        content: str,
        source_path: str
    ) -> List[DocumentChunk]:
        """
        Semantic chunking that preserves document structure.

        Features:
        - Respects section boundaries
        - Accurate token counting with tiktoken
        - Overlap between chunks for context continuity
        - Metadata includes section headers
        """
        chunks = []

        # First, split into semantic sections
        sections = self._split_into_sections(content)

        for section in sections:
            section_content = section["content"]
            section_header = section["header"]

            if not section_content.strip():
                continue

            # If section fits in one chunk, keep it together
            section_tokens = self._count_tokens(section_content)

            if section_tokens <= self.chunk_size:
                # Section fits in one chunk
                chunks.append(DocumentChunk(
                    chunk_id=f"{source_path}_{len(chunks)}",
                    document_id="",
                    chunk_index=len(chunks),
                    content=section_content,
                    token_count=section_tokens,
                    metadata={
                        "source": source_path,
                        "section": section_header,
                        "is_complete_section": True
                    }
                ))
            else:
                # Section too large - split by sentences with overlap
                sentences = self._split_into_sentences(section_content)
                current_chunk = []
                current_tokens = 0

                for sentence in sentences:
                    sentence_tokens = self._count_tokens(sentence)

                    if current_tokens + sentence_tokens > self.chunk_size and current_chunk:
                        # Save current chunk
                        chunk_text = ' '.join(current_chunk)
                        chunks.append(DocumentChunk(
                            chunk_id=f"{source_path}_{len(chunks)}",
                            document_id="",
                            chunk_index=len(chunks),
                            content=chunk_text,
                            token_count=current_tokens,
                            metadata={
                                "source": source_path,
                                "section": section_header,
                                "is_complete_section": False
                            }
                        ))

                        # Start new chunk with overlap (last 2 sentences)
                        overlap = current_chunk[-2:] if len(current_chunk) >= 2 else current_chunk[-1:]
                        current_chunk = overlap + [sentence]
                        current_tokens = sum(self._count_tokens(s) for s in current_chunk)
                    else:
                        current_chunk.append(sentence)
                        current_tokens += sentence_tokens

                # Save final chunk from section
                if current_chunk:
                    chunk_text = ' '.join(current_chunk)
                    chunks.append(DocumentChunk(
                        chunk_id=f"{source_path}_{len(chunks)}",
                        document_id="",
                        chunk_index=len(chunks),
                        content=chunk_text,
                        token_count=self._count_tokens(chunk_text),
                        metadata={
                            "source": source_path,
                            "section": section_header,
                            "is_complete_section": False
                        }
                    ))

        return chunks

    def _split_into_sentences(self, text: str) -> List[str]:
        """Split text into sentences, preserving meaningful boundaries."""
        import re

        # Handle common abbreviations to avoid false splits
        text = re.sub(r'\b(Mr|Mrs|Ms|Dr|Prof|Inc|Ltd|Corp|vs|etc|e\.g|i\.e)\.\s', r'\1<PERIOD> ', text)

        # Split on sentence boundaries
        sentences = re.split(r'(?<=[.!?])\s+', text)

        # Restore periods in abbreviations
        sentences = [s.replace('<PERIOD>', '.') for s in sentences]

        # Filter empty sentences
        return [s.strip() for s in sentences if s.strip()]

    async def _batch_embed(self, texts: List[str]) -> List[List[float]]:
        """Batch embed texts using local model or OpenAI."""
        if USE_LOCAL_EMBEDDINGS:
            try:
                from local_embeddings import embed_batch
                logger.info(
                    f"Using local embeddings for {len(texts)} texts"
                )
                return embed_batch(texts)
            except ImportError:
                logger.warning(
                    "Local embeddings not available, "
                    "falling back to OpenAI"
                )

        client = await self._get_openai_client()

        all_embeddings = []

        # Batch up to 2048 texts per request
        for i in range(0, len(texts), 2048):
            batch = texts[i:i + 2048]

            response = await client.embeddings.create(
                model=self.embedding_model,
                input=batch
            )

            all_embeddings.extend([e.embedding for e in response.data])

        return all_embeddings

    async def _check_duplicate(self, content_hash: str) -> Optional[Dict[str, Any]]:
        """Check if document already exists by content hash."""
        # For now, return None (no duplicate check without database)
        # In production, query the documents table
        return None

    async def _store_document_record(self, document: Document):
        """Store document metadata in database."""
        try:
            from database import SessionLocal, KnowledgeBaseItem
            db = SessionLocal()

            item = KnowledgeBaseItem(
                title=document.filename,
                content_text="",  # Content is stored in vector chunks
                category="document",
                source=document.filename,
                content_type=document.file_type,
                source_type="upload",
                content_hash=document.content_hash,
                created_at=document.created_at
            )

            db.add(item)
            db.commit()
            db.close()

        except Exception as e:
            logger.warning(f"Failed to store document record: {e}")

    async def search(
        self,
        query: str,
        limit: int = 5,
        min_similarity: float = 0.7,
        filter_metadata: Optional[Dict[str, Any]] = None
    ) -> List[SearchResult]:
        """
        Search the knowledge base using semantic similarity.

        Args:
            query: Natural language search query
            limit: Maximum results to return
            min_similarity: Minimum similarity threshold (0.0-1.0)
            filter_metadata: Optional metadata filter (e.g., {"competitor": "Epic"})

        Returns:
            List of SearchResult objects with content and citations
        """
        if not self.vector_store:
            logger.warning("No vector store configured - search unavailable")
            return []

        try:
            # Use vector store's search method (handles embedding internally)
            results = await self.vector_store.search(
                query=query,
                limit=limit,
                min_similarity=min_similarity,
                filter_metadata=filter_metadata
            )

            # Convert to our SearchResult format
            return [
                SearchResult(
                    chunk_id=str(r.chunk_id),
                    document_id=r.document_id,
                    content=r.content,
                    metadata=r.metadata or {},
                    similarity=r.similarity
                )
                for r in results
            ]

        except Exception as e:
            logger.error(f"Knowledge base search failed: {e}")
            return []

    async def get_context_for_query(
        self,
        query: str,
        max_chunks: int = 5,
        max_tokens: int = 4000,
        filter_metadata: Optional[Dict[str, Any]] = None,
        timeout_seconds: float = 10.0
    ) -> Dict[str, Any]:
        """
        Get context for a query to use in RAG with full citations.

        This is the main method agents should use to retrieve relevant context
        from the knowledge base. Returns formatted context with source citations.

        Args:
            query: Natural language query
            max_chunks: Maximum chunks to retrieve (default: 5)
            max_tokens: Maximum total tokens in context (default: 4000)
            filter_metadata: Optional filter (e.g., {"competitor": "Epic"})
            timeout_seconds: Timeout for search operation (default: 10.0)

        Returns:
            Dictionary with:
            - context: Formatted context string with source labels
            - citations: List of citation objects with metadata
            - chunks_used: Number of chunks included
            - total_tokens: Estimated token count
            - query: Original query (for debugging)
            - error: Error message if timeout or failure occurred
        """
        try:
            results = await asyncio.wait_for(
                self.search(
                    query=query,
                    limit=max_chunks,
                    filter_metadata=filter_metadata
                ),
                timeout=timeout_seconds
            )
        except asyncio.TimeoutError:
            logger.warning(f"KB search timeout after {timeout_seconds}s for query: {query[:50]}...")
            return {
                "context": "",
                "citations": [],
                "chunks_used": 0,
                "total_tokens": 0,
                "query": query,
                "error": f"Search timed out after {timeout_seconds} seconds"
            }
        except Exception as e:
            logger.error(f"KB search failed: {e}")
            return {
                "context": "",
                "citations": [],
                "chunks_used": 0,
                "total_tokens": 0,
                "query": query,
                "error": str(e)
            }

        if not results:
            return {
                "context": "",
                "citations": [],
                "chunks_used": 0,
                "total_tokens": 0,
                "query": query,
                "message": "No relevant content found in knowledge base"
            }

        # Build context respecting token limit
        context_parts = []
        citations = []
        total_tokens = 0

        for i, result in enumerate(results):
            # Use accurate token counting if available
            chunk_tokens = self._count_tokens(result.content)

            if total_tokens + chunk_tokens > max_tokens:
                break

            # Format context with clear source attribution
            source_label = f"[Source {i + 1}]"
            section_info = result.metadata.get("section", "")
            if section_info:
                source_label += f" ({section_info})"

            context_parts.append(f"{source_label}\n{result.content}")

            # Build comprehensive citation
            citations.append({
                "source_number": i + 1,
                "document_id": result.document_id,
                "chunk_id": result.chunk_id,
                "section": result.metadata.get("section", "Unknown"),
                "source_file": result.metadata.get("source", "Unknown"),
                "content_preview": result.content[:150] + "..." if len(result.content) > 150 else result.content,
                "similarity_score": round(result.similarity, 3),
                "is_complete_section": result.metadata.get("is_complete_section", False)
            })
            total_tokens += chunk_tokens

        return {
            "context": "\n\n".join(context_parts),
            "citations": citations,
            "chunks_used": len(context_parts),
            "total_tokens": total_tokens,
            "query": query
        }

    async def generate_rag_response(
        self,
        query: str,
        system_prompt: Optional[str] = None,
        max_context_tokens: int = 4000,
        model: str = "gpt-4.1"
    ) -> Dict[str, Any]:
        """
        Complete RAG pipeline: retrieve context and generate response.

        This method combines retrieval and generation for a complete RAG response.

        Args:
            query: User question or request
            system_prompt: Optional custom system prompt
            max_context_tokens: Max tokens for retrieved context
            model: OpenAI model to use for generation

        Returns:
            Dictionary with:
            - response: Generated answer
            - citations: Source citations used
            - model: Model used for generation
            - usage: Token usage statistics
        """
        # Step 1: Retrieve relevant context
        context_result = await self.get_context_for_query(
            query=query,
            max_tokens=max_context_tokens
        )

        if not context_result["context"]:
            return {
                "response": "I don't have enough information in my knowledge base to answer that question.",
                "citations": [],
                "model": model,
                "usage": {"prompt_tokens": 0, "completion_tokens": 0}
            }

        # Step 2: Build the prompt
        default_system = """You are a competitive intelligence analyst for Certify Health.
Answer questions based ONLY on the provided context from our knowledge base.
Always cite your sources using [Source N] notation.
If the context doesn't contain enough information, say so clearly."""

        system = system_prompt or default_system

        user_prompt = f"""Context from knowledge base:
{context_result['context']}

---

Question: {query}

Please answer based on the context above. Cite specific sources using [Source N] notation."""

        # Step 3: Generate response
        client = await self._get_openai_client()

        response = await client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": system},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.3,  # Lower temperature for factual responses
            max_tokens=1000
        )

        return {
            "response": response.choices[0].message.content,
            "citations": context_result["citations"],
            "model": model,
            "usage": {
                "prompt_tokens": response.usage.prompt_tokens,
                "completion_tokens": response.usage.completion_tokens,
                "total_tokens": response.usage.total_tokens
            },
            "chunks_used": context_result["chunks_used"]
        }


# CLI testing
if __name__ == "__main__":

    async def test():
        kb = KnowledgeBase()

        # Test chunking
        test_content = """
        This is paragraph one about competitive intelligence.
        It contains important information about market analysis.

        This is paragraph two about competitor tracking.
        We monitor their pricing and product strategies.

        This is paragraph three about sales enablement.
        Battlecards help sales teams win more deals.
        """

        chunks = kb._chunk_content(test_content, "test.txt")
        print(f"Created {len(chunks)} chunks")
        for chunk in chunks:
            print(f"  Chunk {chunk.chunk_index}: {len(chunk.content)} chars, ~{chunk.token_count} tokens")

    asyncio.run(test())
